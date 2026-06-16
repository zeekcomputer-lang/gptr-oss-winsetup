#!/usr/bin/env python3
"""
md_to_docx — 마크다운 보고서 → 세련된 비즈니스 DOCX 변환 (표 지원)

deep-doc-pipeline-v2 의 scripts/md_to_docx.py 를 참조해, gptr-oss-winsetup 의
출력(research/chrono 보고서 .md)에 맞춰 **개선**한 버전.

deepdoc-v2 대비 개선점:
  - ✅ **GFM 표(table)** 렌더링 — 헤더 음영·테두리·교대 행 음영·셀 인라인 서식
  - ✅ **순서 목록**(1. 2. 3.) + 중첩 들여쓰기 불릿/번호
  - ✅ **펜스 코드블록**(``` ```) — 모노스페이스 음영 블록
  - ✅ **링크** [text](url) — 실제 하이퍼링크(밑줄·강조색)
  - ✅ **H4** 지원 + 본문 인라인(굵게/코드/링크) 통합 파서
  - ✅ 표지(H1) + 머리말 메타 + 자동 페이지 번호 (deepdoc 스타일 유지)

사용:
    python tools/md_to_docx.py outputs/report.md
    python tools/md_to_docx.py outputs/report.md -o report.docx
    python tools/md_to_docx.py a.md b.md -o combined.docx   # 순서대로 병합(페이지 구분)

의존성:  pip install python-docx   (setup.py 가 자동 설치)
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 미설치. 설치 후 재실행:")
    print("  pip install python-docx        (또는 python tools/setup.py)")
    sys.exit(1)


# ──────────────────────────────────────────────────────────────
# 스타일 상수 (비즈니스 보고서 팔레트 — deepdoc-v2 계승)
# ──────────────────────────────────────────────────────────────
FONT_NAME = "맑은 고딕"
FONT_MONO = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_H1 = Pt(24)
FONT_SIZE_H2 = Pt(15)
FONT_SIZE_H3 = Pt(12)
FONT_SIZE_H4 = Pt(11)
LINE_SPACING = Pt(18)
PAGE_MARGIN = Cm(2.4)

COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy — 제목/헤딩
COLOR_ACCENT = RGBColor(0x2E, 0x6D, 0xB4)    # blue — 강조/링크
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)      # near-black 본문
COLOR_MUTED = RGBColor(0x70, 0x70, 0x70)     # gray — 메타
COLOR_RULE = "1F3A5F"                         # 헤딩 밑줄
COLOR_TABLE_HEADER_BG = "1F3A5F"             # 표 헤더 음영(네이비)
COLOR_TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_ROW_ALT = "EEF3FA"               # 교대 행 음영(연파랑)
COLOR_TABLE_BORDER = "C9D4E2"                # 표 테두리(연회색)
COLOR_CODE_BG = "F4F5F7"                       # 코드블록 음영


# ──────────────────────────────────────────────────────────────
# 정규식
# ──────────────────────────────────────────────────────────────
_RE_H1 = re.compile(r'^# (.+)$')
_RE_H2 = re.compile(r'^## (.+)$')
_RE_H3 = re.compile(r'^### (.+)$')
_RE_H4 = re.compile(r'^#### (.+)$')
_RE_BULLET = re.compile(r'^(\s*)[-*+] (.+)$')
_RE_ORDERED = re.compile(r'^(\s*)\d+\. (.+)$')
_RE_BLOCKQUOTE = re.compile(r'^> ?(.*)$')
_RE_HR = re.compile(r'^(-{3,}|\*{3,}|_{3,})$')
_RE_FENCE = re.compile(r'^```(.*)$')
_RE_TABLE_ROW = re.compile(r'^\s*\|.*\|\s*$')
_RE_TABLE_SEP = re.compile(r'^\s*\|?[\s:\-|]+\|?\s*$')
# 인라인: **bold**, `code`, [text](url)
_RE_INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
_RE_BOLD = re.compile(r'\*\*(.+?)\*\*')
_RE_CODE = re.compile(r'`([^`]+)`')
_RE_LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


# ──────────────────────────────────────────────────────────────
# 저수준 XML 헬퍼
# ──────────────────────────────────────────────────────────────
def _set_cn_font(run, name: str = FONT_NAME):
    """East-Asian 폰트 적용(python-docx 가 기본으로 w:eastAsia 를 안 넣음)."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)


def _shade(element_pr, fill_hex: str):
    """pPr/tcPr 에 음영(w:shd) 적용."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    element_pr.append(shd)


def _shade_paragraph(paragraph, fill_hex: str):
    _shade(paragraph._p.get_or_add_pPr(), fill_hex)


def _bottom_border(paragraph, color_hex: str, size: int = 12):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _add_hyperlink(paragraph, url: str, text: str):
    """실제 클릭 가능한 하이퍼링크 run 추가(밑줄+강조색)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E6DB4')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), FONT_NAME)
    rpr.append(rfonts)
    rpr.append(color)
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ──────────────────────────────────────────────────────────────
# 인라인 서식 파서 (bold / code / link)
# ──────────────────────────────────────────────────────────────
def _add_styled_runs(paragraph, text: str, base_color: RGBColor | None = None,
                     base_size: Pt | None = None):
    parts = _RE_INLINE.split(text)
    for part in parts:
        if not part:
            continue
        m_bold = _RE_BOLD.fullmatch(part)
        m_code = _RE_CODE.fullmatch(part)
        m_link = _RE_LINK.fullmatch(part)
        if m_bold:
            run = paragraph.add_run(m_bold.group(1))
            run.bold = True
            _set_cn_font(run)
            if base_color:
                run.font.color.rgb = base_color
            if base_size:
                run.font.size = base_size
        elif m_code:
            run = paragraph.add_run(m_code.group(1))
            run.font.name = FONT_MONO
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif m_link:
            _add_hyperlink(paragraph, m_link.group(2), m_link.group(1))
        else:
            run = paragraph.add_run(part)
            _set_cn_font(run)
            if base_color:
                run.font.color.rgb = base_color
            if base_size:
                run.font.size = base_size


# ──────────────────────────────────────────────────────────────
# 표(table) 렌더링 — 핵심 개선
# ──────────────────────────────────────────────────────────────
def _split_table_row(line: str) -> list[str]:
    """| a | b | c | → ['a','b','c'] (이스케이프 \\| 보존)."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    # 이스케이프된 파이프 \| 임시치환
    cells = re.split(r'(?<!\\)\|', s)
    return [c.replace(r'\|', '|').strip() for c in cells]


def _set_cell_border(cell, color_hex: str = COLOR_TABLE_BORDER, sz: int = 4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        borders.append(el)
    tc_pr.append(borders)


def _style_cell(cell, text: str, *, header: bool, alt: bool, align: str = "left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if header:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TABLE_HEADER_FG
        _set_cn_font(run)
        _shade(cell._tc.get_or_add_tcPr(), COLOR_TABLE_HEADER_BG)
    else:
        _add_styled_runs(p, text, base_size=Pt(10))
        if alt:
            _shade(cell._tc.get_or_add_tcPr(), COLOR_TABLE_ROW_ALT)
    _set_cell_border(cell)


def _parse_alignments(sep_line: str, ncols: int) -> list[str]:
    aligns = []
    for c in _split_table_row(sep_line):
        c = c.strip()
        left = c.startswith(":")
        right = c.endswith(":")
        aligns.append("center" if left and right else "right" if right else "left")
    while len(aligns) < ncols:
        aligns.append("left")
    return aligns[:ncols]


def _add_table(doc: Document, header: list[str], rows: list[list[str]], aligns: list[str]):
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # 헤더
    for j, h in enumerate(header):
        _style_cell(table.rows[0].cells[j], h, header=True, alt=False,
                    align=aligns[j] if j < len(aligns) else "left")
    # 본문
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            _style_cell(cells[j], val, header=False, alt=(i % 2 == 1),
                        align=aligns[j] if j < len(aligns) else "left")
    # 표 뒤 간격
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return table


# ──────────────────────────────────────────────────────────────
# 코드블록
# ──────────────────────────────────────────────────────────────
def _add_code_block(doc: Document, code_lines: list[str], lang: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _shade_paragraph(p, COLOR_CODE_BG)
    text = "\n".join(code_lines)
    run = p.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)


# ──────────────────────────────────────────────────────────────
# 문서 스타일/표지/페이지번호 (deepdoc-v2 계승)
# ──────────────────────────────────────────────────────────────
def _setup_doc_styles(doc: Document):
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = Pt(6)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT_NAME)


def _add_cover_title(doc: Document, title: str, label: str = "BUSINESS REPORT"):
    lab = doc.add_paragraph()
    lab.paragraph_format.space_after = Pt(2)
    lr = lab.add_run(label)
    lr.font.size = Pt(9)
    lr.bold = True
    lr.font.color.rgb = COLOR_ACCENT
    _set_cn_font(lr)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    _set_cn_font(run)
    run.font.size = FONT_SIZE_H1
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    _bottom_border(p, COLOR_RULE, size=18)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_page_numbers(doc: Document):
    def _field(instr: str):
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), instr)
        return fld
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        run = p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED
        p._p.append(_field('PAGE'))
        sep = p.add_run(" / ")
        sep.font.size = Pt(8)
        sep.font.color.rgb = COLOR_MUTED
        p._p.append(_field('NUMPAGES'))


# ──────────────────────────────────────────────────────────────
# 메인 변환기
# ──────────────────────────────────────────────────────────────
def md_to_docx(md_text: str, doc: Document, title_label: str = "",
               cover_label: str = "BUSINESS REPORT"):
    if title_label:
        p = doc.add_paragraph()
        run = p.add_run(title_label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lines = md_text.split("\n")
    i, n = 0, len(lines)
    bq_buf: list[str] = []

    def _flush_bq():
        if not bq_buf:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _bottom_border(p, "C9D4E2", size=4)  # 가벼운 구분 느낌
        _add_styled_runs(p, " ".join(bq_buf), base_color=RGBColor(0x66, 0x66, 0x66))
        for r in p.runs:
            r.font.size = Pt(9.5)
        bq_buf.clear()

    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # 빈 줄
        if not stripped:
            _flush_bq()
            i += 1
            continue

        # 펜스 코드블록
        mfence = _RE_FENCE.match(stripped)
        if mfence:
            _flush_bq()
            lang = mfence.group(1).strip()
            code: list[str] = []
            i += 1
            while i < n and not _RE_FENCE.match(lines[i].strip()):
                code.append(lines[i])
                i += 1
            i += 1  # 닫는 ``` 소비
            _add_code_block(doc, code, lang)
            continue

        # 표: 현재 줄이 표행 + 다음 줄이 구분행이면 표 블록 수집
        if _RE_TABLE_ROW.match(raw) and i + 1 < n and _RE_TABLE_SEP.match(lines[i + 1]) \
                and "-" in lines[i + 1]:
            _flush_bq()
            header = _split_table_row(raw)
            aligns = _parse_alignments(lines[i + 1], len(header))
            i += 2
            body: list[list[str]] = []
            while i < n and _RE_TABLE_ROW.match(lines[i]):
                body.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, header, body, aligns)
            continue

        # 수평선
        if _RE_HR.match(stripped):
            _flush_bq()
            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(6)
            rule.paragraph_format.space_after = Pt(6)
            _bottom_border(rule, "C9D4E2", size=6)
            i += 1
            continue

        # 인용
        mbq = _RE_BLOCKQUOTE.match(stripped) if stripped.startswith(">") else None
        if mbq:
            bq_buf.append(mbq.group(1))
            i += 1
            continue
        else:
            _flush_bq()

        # 헤딩
        m = _RE_H1.match(stripped)
        if m:
            _add_cover_title(doc, m.group(1), cover_label)
            i += 1
            continue
        m = _RE_H2.match(stripped)
        if m:
            p = doc.add_heading("", level=2)
            run = p.add_run(m.group(1))
            _set_cn_font(run)
            run.font.size = FONT_SIZE_H2
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            _bottom_border(p, COLOR_RULE, size=8)
            i += 1
            continue
        m = _RE_H3.match(stripped)
        if m:
            p = doc.add_heading("", level=3)
            run = p.add_run(m.group(1))
            _set_cn_font(run)
            run.font.size = FONT_SIZE_H3
            run.bold = True
            run.font.color.rgb = COLOR_ACCENT
            i += 1
            continue
        m = _RE_H4.match(stripped)
        if m:
            p = doc.add_heading("", level=4)
            run = p.add_run(m.group(1))
            _set_cn_font(run)
            run.font.size = FONT_SIZE_H4
            run.bold = True
            run.font.color.rgb = COLOR_ACCENT
            i += 1
            continue

        # 순서 목록 (들여쓰기 → level)
        mo = _RE_ORDERED.match(raw)
        if mo:
            indent = len(mo.group(1).replace("\t", "  "))
            lvl = min(indent // 2, 2)
            style = "List Number" if lvl == 0 else f"List Number {lvl + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Number")
            _add_styled_runs(p, mo.group(2))
            i += 1
            continue

        # 불릿 목록 (들여쓰기 → level)
        mb = _RE_BULLET.match(raw)
        if mb:
            indent = len(mb.group(1).replace("\t", "  "))
            lvl = min(indent // 2, 2)
            style = "List Bullet" if lvl == 0 else f"List Bullet {lvl + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_styled_runs(p, mb.group(2))
            i += 1
            continue

        # 일반 단락
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_styled_runs(p, stripped)
        i += 1

    _flush_bq()


def build_report_docx(md_text: str, out_path, cover_label: str = "BUSINESS REPORT"):
    """단일 마크다운 보고서 → 비즈니스 DOCX. run_research 자동 내보내기에서 호출."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _setup_doc_styles(doc)
    md_to_docx(md_text, doc, title_label="", cover_label=cover_label)
    _add_page_numbers(doc)
    doc.save(str(out_path))
    return out_path


# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────
def main() -> int:
    ap = argparse.ArgumentParser(
        description="마크다운 보고서 → 비즈니스 DOCX(표 지원)",
        epilog="복수 파일 전달 시 순서대로 병합(페이지 구분).",
    )
    ap.add_argument("inputs", nargs="+", help="입력 마크다운 파일(1개 이상)")
    ap.add_argument("-o", "--output", default=None,
                    help="출력 DOCX 경로(기본: 첫 입력파일명.docx)")
    ap.add_argument("--cover-label", default="BUSINESS REPORT",
                    help="표지 상단 라벨(기본 BUSINESS REPORT)")
    args = ap.parse_args()

    input_paths = []
    for inp in args.inputs:
        p = Path(inp)
        if not p.exists():
            print(f"[ERROR] 파일 없음: {p}")
            return 1
        input_paths.append(p)

    out_path = Path(args.output) if args.output else input_paths[0].with_suffix(".docx")

    doc = Document()
    _setup_doc_styles(doc)
    for idx, inp in enumerate(input_paths):
        if idx > 0:
            doc.add_page_break()
        md_text = inp.read_text(encoding="utf-8")
        label = inp.name if len(input_paths) > 1 else ""
        md_to_docx(md_text, doc, title_label=label, cover_label=args.cover_label)
        print(f"  [{idx + 1}/{len(input_paths)}] {inp.name} ({len(md_text):,} chars)")
    _add_page_numbers(doc)
    doc.save(str(out_path))
    print(f"✅ 저장: {out_path.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
